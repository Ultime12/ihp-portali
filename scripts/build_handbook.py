from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "IHP-Sistem-Kullanim-Kitapcigi.md"
OUTPUT = ROOT / "output" / "docx" / "IHP-Sistem-Kullanim-Kitapcigi.docx"
LOGO = ROOT / "assets" / "identity" / "party-mark.png"

NAVY = "071426"
NAVY_2 = "0E2140"
BLUE = "4F8CFF"
VIOLET = "8B7CFF"
TEXT = "1D2939"
MUTED = "536273"
LIGHT = "EEF4FF"
LINE = "D8E2F0"
WHITE = "FFFFFF"
GOLD = "D7B56D"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=LINE, size="6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph, keep=True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    node = p_pr.find(qn("w:keepNext"))
    if keep and node is None:
        node = OxmlElement("w:keepNext")
        p_pr.append(node)
    elif not keep and node is not None:
        p_pr.remove(node)


def set_keep_lines(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:keepLines")) is None:
        p_pr.append(OxmlElement("w:keepLines"))


def set_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("SAYFA ")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.25

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string(BLUE)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor.from_string(BLUE)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)

    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor.from_string("1F4D78")
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(5)

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor.from_string(TEXT)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def configure_section(section, cover=False) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.5 if cover else 1.0)
    section.bottom_margin = Inches(0.5 if cover else 0.78)
    section.left_margin = Inches(0.5 if cover else 1.0)
    section.right_margin = Inches(0.5 if cover else 1.0)
    section.header_distance = Inches(0.3 if cover else 0.492)
    section.footer_distance = Inches(0.3 if cover else 0.42)


def add_inline_runs(paragraph, text: str, color: str | None = None) -> None:
    parts = re.split(r"(\*\*.+?\*\*|https?://\S+)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
            if part.startswith("http://") or part.startswith("https://"):
                run.font.color.rgb = RGBColor.from_string(BLUE)
                run.underline = True
        if color:
            run.font.color.rgb = RGBColor.from_string(color)


def add_cover(doc: Document) -> None:
    section = doc.sections[0]
    configure_section(section, cover=True)
    section.different_first_page_header_footer = True

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(7.45)
    cell = table.cell(0, 0)
    cell.width = Inches(7.45)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, NAVY)
    set_cell_margins(cell, top=500, start=520, bottom=480, end=520)
    set_table_borders(table, color=NAVY, size="0")

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO.exists():
        p.add_run().add_picture(str(LOGO), width=Inches(1.42))

    kicker = cell.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_before = Pt(20)
    kicker.paragraph_format.space_after = Pt(12)
    run = kicker.add_run("İHP  •  DİJİTAL SİSTEMLER")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(GOLD)

    title = cell.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("KULLANIM")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(30)
    run.font.color.rgb = RGBColor.from_string(WHITE)
    title.add_run("\n")
    run = title.add_run("KİTAPÇIĞI")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(30)
    run.font.color.rgb = RGBColor.from_string(VIOLET)

    subtitle = cell.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(8)
    subtitle.paragraph_format.space_after = Pt(26)
    run = subtitle.add_run("Ana Portal  •  Disiplin Kurulu  •  Finans  •  Kurumsal Posta")
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string("C8D5EA")

    line = cell.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line.paragraph_format.space_after = Pt(24)
    run = line.add_run("━━━━━━━━━━━━━━━━━━━━")
    run.font.color.rgb = RGBColor.from_string(BLUE)

    meta = cell.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(38)
    run = meta.add_run("19 TEMMUZ 2026\nÖĞRENCİ TOPLULUĞU DİJİTAL ÇALIŞMA ALANLARI")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("9EB3D0")

    doc.add_page_break()


def add_header_footer(section) -> None:
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("İHP DİJİTAL SİSTEMLER")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(BLUE)
    run = p.add_run("   •   KULLANIM KİTAPÇIĞI")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), LINE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    set_page_number(footer_p)


def add_contents(doc: Document, headings: list[str]) -> None:
    title = doc.add_heading("İçindekiler", level=1)
    title.paragraph_format.space_before = Pt(0)
    intro = doc.add_paragraph("Uygulama adına veya işlem başlığına göre ilgili bölüme gidin.")
    intro.paragraph_format.space_after = Pt(12)
    for index, heading in enumerate(headings, start=1):
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.columns[0].width = Inches(0.45)
        table.columns[1].width = Inches(5.92)
        set_table_borders(table, color=WHITE, size="0")
        num_cell, text_cell = table.rows[0].cells
        set_cell_shading(num_cell, NAVY_2)
        set_cell_shading(text_cell, "F6F8FC" if index % 2 else LIGHT)
        set_cell_margins(num_cell, 75, 85, 75, 85)
        set_cell_margins(text_cell, 75, 120, 75, 120)
        p_num = num_cell.paragraphs[0]
        p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_num.add_run(f"{index:02d}")
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(WHITE)
        p_text = text_cell.paragraphs[0]
        run = p_text.add_run(re.sub(r"^\d+\.\s*", "", heading))
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string(TEXT)
    doc.add_page_break()


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(0.12)
    table.columns[1].width = Inches(6.15)
    set_table_borders(table, color=WHITE, size="0")
    accent, body = table.rows[0].cells
    set_cell_shading(accent, VIOLET)
    set_cell_shading(body, LIGHT)
    set_cell_margins(accent, 70, 0, 70, 0)
    set_cell_margins(body, 110, 180, 110, 180)
    p = body.paragraphs[0]
    add_inline_runs(p, text, color=TEXT)
    p.paragraph_format.space_after = Pt(0)


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if len(rows) < 2:
        return
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    widths = [1.22, 1.65, 3.5] if len(rows[0]) == 3 else [6.32 / len(rows[0])] * len(rows[0])
    for idx, width in enumerate(widths):
        table.columns[idx].width = Inches(width)
    for idx, value in enumerate(rows[0]):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, NAVY_2)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        add_inline_runs(p, value.strip(), color=WHITE)
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(9)
    set_repeat_table_header(table.rows[0])
    for row_index, row in enumerate(rows[1:]):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            set_cell_shading(cell, "F8FAFD" if row_index % 2 == 0 else WHITE)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            add_inline_runs(p, value.strip())
            for run in p.runs:
                run.font.size = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        values = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", value) for value in values):
            rows.append(values)
        index += 1
    return rows, index


def add_body(doc: Document, lines: list[str]) -> None:
    index = 0
    first_major = True
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        text = " ".join(part.strip() for part in paragraph_buffer).strip()
        if text:
            p = doc.add_paragraph()
            add_inline_runs(p, text)
        paragraph_buffer = []

    while index < len(lines):
        raw = lines[index].rstrip()
        stripped = raw.strip()

        if not stripped:
            flush_paragraph()
            index += 1
            continue

        if stripped.startswith("|"):
            flush_paragraph()
            rows, index = parse_table(lines, index)
            add_markdown_table(doc, rows)
            continue

        heading_match = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading_match:
            flush_paragraph()
            marks, title = heading_match.groups()
            level = len(marks) - 1
            if level == 1:
                if not first_major:
                    doc.add_page_break()
                first_major = False
            p = doc.add_heading(title, level=min(level, 3))
            set_keep_with_next(p)
            set_keep_lines(p)
            index += 1
            continue

        if stripped.startswith("> "):
            flush_paragraph()
            add_callout(doc, stripped[2:].strip())
            index += 1
            continue

        bullet_match = re.match(r"^-\s+(.+)$", stripped)
        if bullet_match:
            flush_paragraph()
            value = bullet_match.group(1)
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, value)
            set_keep_lines(p)
            index += 1
            continue

        number_match = re.match(r"^\d+\.\s+(.+)$", stripped)
        if number_match:
            flush_paragraph()
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, number_match.group(1))
            set_keep_lines(p)
            index += 1
            continue

        paragraph_buffer.append(stripped)
        index += 1

    flush_paragraph()


def remove_trailing_empty_paragraphs(doc: Document) -> None:
    for paragraph in reversed(doc.paragraphs):
        if paragraph.text.strip():
            break
        paragraph._element.getparent().remove(paragraph._element)


def build() -> None:
    markdown = SOURCE.read_text(encoding="utf-8")
    lines = markdown.splitlines()
    body_start = next(index for index, line in enumerate(lines) if line.startswith("## 1."))
    body_lines = lines[body_start:]
    headings = [line[3:].strip() for line in body_lines if line.startswith("## ")]

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    add_cover(doc)

    body_section = doc.add_section(WD_SECTION.CONTINUOUS)
    configure_section(body_section)
    add_header_footer(body_section)

    add_contents(doc, headings)
    add_callout(
        doc,
        "Bu kitapçık işlemlerin nerede ve nasıl yapıldığını anlatır. Yetki, süre ve karar konularında sistemde yayımlanan güncel yönetmelik PDF'leri esastır.",
    )
    doc.add_paragraph()
    add_body(doc, body_lines)
    remove_trailing_empty_paragraphs(doc)

    core = doc.core_properties
    core.title = "İHP Dijital Sistemler Kullanım Kitapçığı"
    core.subject = "Ana Portal, Disiplin Kurulu, Finans ve Kurumsal Posta kullanım rehberi"
    core.author = "İHP Öğrenci Topluluğu"
    core.keywords = "İHP, portal, disiplin kurulu, finans, mail, kullanım kitapçığı"
    core.comments = "Sistem menüleri ve canlı işlem akışları esas alınarak hazırlanmıştır."

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
